"""作業報告書自動生成 Streamlit クライアント.

作業写真、音声メモ、作業内容メモ、設備型番を入力し、
作業報告書自動生成エージェント（/generate-work-report）を呼び出して
標準フォーマット作業報告書と請求項目候補を表示するアプリです。

環境変数:
    API_BASE_URL: エージェントのベースURL（例: https://agent.example.com/xxxx）
    API_KEY: x-api-key に設定するAPIキー
    API_TIMEOUT_SEC: タイムアウト秒数（任意。既定値 600）
"""

from __future__ import annotations

import html
import io
import json
import mimetypes
import os
import sys
import traceback
from datetime import date
from typing import Any, Dict, Iterable, List, Optional, Tuple

import requests
import streamlit as st

DOCX_IMPORT_ERROR: str = ""
try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except Exception as import_error:
    Document = None
    Pt = None
    DOCX_AVAILABLE = False
    DOCX_IMPORT_ERROR = repr(import_error)



try:
    from streamlit_back_camera_input import back_camera_input
except Exception:  # pragma: no cover - optional dependency fallback
    back_camera_input = None


# -----------------------------------------------------------
# 設定
# -----------------------------------------------------------
API_BASE_URL: str = os.environ.get("API_BASE_URL", "").strip()
API_KEY: str = os.environ.get("API_KEY", "").strip()


def parse_timeout(default: int = 600) -> int:
    """環境変数 API_TIMEOUT_SEC を秒として読み取る。"""
    raw = os.environ.get("API_TIMEOUT_SEC", "").strip()
    if not raw:
        return default
    try:
        value = int(raw)
    except ValueError:
        return default
    return value if value > 0 else default


TIMEOUT_SEC: int = parse_timeout(600)

EQUIPMENT_OPTIONS: List[str] = [
    "CV-3200",
    "PR-500A",
    "OV-210",
    "PK-700",
    "CP-100",
]



st.set_page_config(
    page_title="作業報告書自動生成",
    layout="wide",
    initial_sidebar_state="collapsed",
)


# -----------------------------------------------------------
# CSS
# -----------------------------------------------------------
_CSS = """
<style>
html, body, [class*="css"] {
  font-family: 'Yu Gothic','YuGothic','Hiragino Kaku Gothic ProN','Meiryo',sans-serif;
  color:#1a1a1a;
}
.main .block-container { padding-top: 1.4rem; }
.hero-banner {
  background:#000;border-radius:14px;padding:28px 36px;margin-bottom:22px;color:#fff;
  position:relative;overflow:hidden;
}
.hero-banner::before {
  content:'';position:absolute;top:-30%;right:-10%;width:54%;height:170%;
  background:linear-gradient(135deg,#575B7C 0%,#4455E4 100%);
  clip-path:polygon(60% 0%,100% 0%,100% 100%,18% 100%);opacity:.20;
}
.hero-banner h1 { margin:0 0 6px;font-size:1.65rem;font-weight:800;position:relative; }
.hero-banner p { margin:0;font-size:.92rem;opacity:.78;position:relative;line-height:1.7; }
.section-card {
  background:#fff;border:1px solid #E8E8EC;border-radius:12px;padding:22px;margin-bottom:16px;
  box-shadow:0 2px 8px rgba(0,0,0,.04);
}
.section-title {
  font-size:1.05rem;font-weight:800;color:#000;margin:0 0 16px;display:flex;align-items:center;gap:8px;
  border-bottom:2px solid #4455E4;padding-bottom:10px;
}
.icon { display:inline-flex;align-items:center;justify-content:center;width:30px;height:30px;border-radius:8px;background:#E8E8EC; }
.summary-card {
  background:#f9f9fb;border:1px solid #E8E8EC;border-left:4px solid #4455E4;border-radius:10px;
  padding:18px 22px;margin-bottom:16px;
}
.summary-card ul { margin:0;padding-left:18px;line-height:1.8; }
.kv-card {
  background:#fff;border:1px solid #E8E8EC;border-radius:10px;padding:16px 18px;margin-bottom:10px;
}
.kv-label { color:#575B7C;font-size:.78rem;font-weight:800;letter-spacing:.04em;margin-bottom:5px; }
.kv-value { color:#111;font-size:.95rem;line-height:1.65;white-space:pre-wrap; }
.status-badge {
  display:inline-flex;align-items:center;padding:8px 13px;border-radius:999px;font-size:.82rem;font-weight:800;margin:2px 4px 8px 0;
  background:#eef2ff;color:#3730a3;border:1px solid #c7d2fe;
}
.status-ok { background:#ecfdf5;color:#065f46;border-color:#a7f3d0; }
.status-ng { background:#fef2f2;color:#991b1b;border-color:#fecaca; }
.status-warn { background:#fffbeb;color:#92400e;border-color:#fde68a; }
.step-card {
  border:1px solid #E8E8EC;border-radius:10px;overflow:hidden;background:#fff;margin-bottom:12px;
  box-shadow:0 1px 4px rgba(0,0,0,.03);
}
.step-header { background:#000;color:#fff;padding:10px 15px;font-weight:800;font-size:.86rem; }
.step-body { padding:14px 16px; }
.step-body p { margin:.25rem 0 .55rem;line-height:1.65; }
.chip-grid { display:flex;flex-wrap:wrap;gap:10px;margin-top:8px; }
.chip {
  display:inline-flex;align-items:center;gap:8px;padding:9px 14px;border:1px solid #E8E8EC;border-radius:9px;background:#fff;
  font-size:.88rem;font-weight:600;
}
.chip-num { width:24px;height:24px;border-radius:7px;background:#4455E4;color:#fff;display:inline-flex;align-items:center;justify-content:center;font-size:.75rem; }
.metric-row { display:flex;gap:12px;flex-wrap:wrap;margin-bottom:10px; }
.metric-card {
  background:#eef2ff;border:1px solid #c7d2fe;border-radius:10px;padding:18px 20px;min-width:160px;text-align:center;
}
.metric-value { color:#4455E4;font-size:2rem;font-weight:900;line-height:1; }
.metric-label { color:#575B7C;font-size:.78rem;font-weight:800;margin-top:8px; }
.empty-state { text-align:center;padding:42px 20px;color:#8186A5; }
.empty-state .empty-icon { font-size:2.5rem;margin-bottom:10px;opacity:.45; }
.stTabs [data-baseweb="tab-list"] { gap:0;background:#E8E8EC;border-radius:8px;padding:3px; }
.stTabs [data-baseweb="tab"] { border-radius:6px;font-weight:700;font-size:.84rem;padding:8px 12px;color:#575B7C; }
.stTabs [aria-selected="true"] { background:#fff !important;color:#4455E4 !important;box-shadow:0 1px 3px rgba(0,0,0,.08); }
.stTabs [data-baseweb="tab-highlight"], .stTabs [data-baseweb="tab-border"] { display:none; }
.stButton > button[kind="primary"] {
  background:#4455E4 !important;border:none !important;border-radius:8px !important;font-weight:800 !important;
  box-shadow:0 2px 8px rgba(68,85,228,.25) !important;
}

.report-info-table {
  width:100%;
  border-collapse:separate;
  border-spacing:0;
  border:1px solid #E8E8EC;
  border-radius:12px;
  overflow:hidden;
  background:#fff;
  margin-bottom:14px;
  box-shadow:0 1px 4px rgba(0,0,0,.03);
}
.report-info-table th {
  width:18%;
  background:#f4f5fb;
  color:#575B7C;
  font-size:.78rem;
  font-weight:800;
  letter-spacing:.04em;
  text-align:left;
  vertical-align:top;
  padding:12px 14px;
  border-right:1px solid #E8E8EC;
  border-bottom:1px solid #E8E8EC;
  white-space:nowrap;
}
.report-info-table td {
  width:32%;
  color:#111;
  font-size:.94rem;
  line-height:1.65;
  vertical-align:top;
  padding:12px 14px;
  border-right:1px solid #E8E8EC;
  border-bottom:1px solid #E8E8EC;
  white-space:pre-wrap;
  word-break:break-word;
}
.report-info-table tr:last-child th,
.report-info-table tr:last-child td {
  border-bottom:none;
}
.report-info-table th:nth-child(3),
.report-info-table td:nth-child(4) {
  border-right:none;
}
.report-text-table {
  width:100%;
  border-collapse:separate;
  border-spacing:0;
  border:1px solid #E8E8EC;
  border-radius:12px;
  overflow:hidden;
  background:#fff;
  margin-bottom:14px;
  box-shadow:0 1px 4px rgba(0,0,0,.03);
}
.report-text-table th {
  width:24%;
  background:#f4f5fb;
  color:#575B7C;
  font-size:.78rem;
  font-weight:800;
  letter-spacing:.04em;
  text-align:left;
  vertical-align:top;
  padding:13px 15px;
  border-right:1px solid #E8E8EC;
  border-bottom:1px solid #E8E8EC;
  white-space:nowrap;
}
.report-text-table td {
  color:#111;
  font-size:.94rem;
  line-height:1.75;
  vertical-align:top;
  padding:13px 15px;
  border-bottom:1px solid #E8E8EC;
  white-space:pre-wrap;
  word-break:break-word;
}
.report-text-table tr:last-child th,
.report-text-table tr:last-child td {
  border-bottom:none;
}
.report-section-heading {
  font-size:.95rem;
  font-weight:800;
  color:#111;
  margin:12px 0 8px;
  display:flex;
  align-items:center;
  gap:6px;
}

</style>
"""
st.markdown(_CSS, unsafe_allow_html=True)


# -----------------------------------------------------------
# ユーティリティ
# -----------------------------------------------------------
def esc(value: Any) -> str:
    """HTML表示用に安全な文字列へ変換する。"""
    if value is None:
        return ""
    return html.escape(str(value), quote=True)


def safe_dict(value: Any) -> Dict[str, Any]:
    return value if isinstance(value, dict) else {}


def safe_list(value: Any) -> List[Any]:
    return value if isinstance(value, list) else []


def log_stdout(obj: Any, prefix: str = "") -> None:
    """デバッグ情報を標準出力へ出す。画面には出さない。"""
    try:
        text = json.dumps(obj, ensure_ascii=False, indent=2) if isinstance(obj, (dict, list)) else str(obj)
    except Exception:
        text = str(obj)
    if prefix:
        print(prefix, flush=True)
    print(text, flush=True)


def normalize_result(result: Dict[str, Any]) -> Dict[str, Any]:
    """APIレスポンスのラッパー有無を吸収する。references のみ正式キーとして扱う。"""
    if not isinstance(result, dict):
        return {}

    response_obj = result.get("response")
    normalized = response_obj if isinstance(response_obj, dict) else result

    if not isinstance(normalized, dict):
        return {}

    if "references" not in normalized:
        normalized["references"] = []

    return normalized


def get_references(result: Dict[str, Any]) -> List[Any]:
    """正式キー references からデータソース参照配列を取得する。"""
    if not isinstance(result, dict):
        return []
    return safe_list(result.get("references"))


def validate_result_schema(result: Dict[str, Any]) -> Tuple[bool, List[str]]:
    """旧仕様・新仕様の両方を許容して最低限のスキーマ確認を行う。"""
    required_base_keys = [
        "endpoint",
        "report",
        "billing_candidates",
        "missing_information",
        "attached_file_observations",
        "uncertain_transcript_segments",
    ]
    if not isinstance(result, dict):
        return False, ["<root is not object>"]

    missing = [key for key in required_base_keys if key not in result]
    if "references" not in result:
        missing.append("references")

    return len(missing) == 0, missing


def log_schema_mismatch(result: Any, missing_keys: List[str]) -> None:
    """スキーマ不一致時の詳細を標準出力へ必ず出力する。"""
    actual_keys = list(result.keys()) if isinstance(result, dict) else []
    log_stdout(
        {
            "message": "応答スキーマが想定と異なります",
            "missing_keys": missing_keys,
            "actual_keys": actual_keys,
            "required_reference_key": "references",
            "raw_result": result,
        },
        prefix="--- SCHEMA MISMATCH DEBUG ---",
    )


def mime_for_file(name: str, supplied_type: Optional[str]) -> str:
    return supplied_type or mimetypes.guess_type(name)[0] or "application/octet-stream"


def uploaded_to_file_tuple(uploaded_file: Any) -> Tuple[str, Tuple[str, io.BytesIO, str]]:
    """requests の files に渡すタプルへ変換する。"""
    name = getattr(uploaded_file, "name", "uploaded_file")
    supplied_type = getattr(uploaded_file, "type", None)
    try:
        uploaded_file.seek(0)
    except Exception:
        pass
    data = uploaded_file.read()
    return "file", (name, io.BytesIO(data), mime_for_file(name, supplied_type))


def camera_bytes_to_file_tuple(data: bytes) -> Tuple[str, Tuple[str, io.BytesIO, str]]:
    return "file", ("camera_photo.jpg", io.BytesIO(data), "image/jpeg")



def recorded_audio_to_file_tuple(recorded_audio: Any) -> Tuple[str, Tuple[str, io.BytesIO, str]]:
    """マイク録音データを WAV ファイルとして requests の files に渡す形式へ変換する。"""
    try:
        recorded_audio.seek(0)
    except Exception:
        pass
    audio_bytes = recorded_audio.read()
    return "file", ("recorded_voice_memo.wav", io.BytesIO(audio_bytes), "audio/wav")


# -----------------------------------------------------------
# エージェント呼び出し
# -----------------------------------------------------------
def call_api(
    equipment_model: str,
    work_memo: str,
    work_date: str,
    customer_name: str,
    worker_name: str,
    uploaded_files: Iterable[Any],
    captured_photo: Optional[bytes],
    recorded_audio: Optional[Any],
) -> Dict[str, Any]:
    """作業報告書自動生成エージェントを multipart/form-data で呼び出す。"""
    endpoint = f"{API_BASE_URL.rstrip('/')}/generate-work-report"

    data = {
        "equipment_model": equipment_model,
        "work_memo": work_memo,
        "work_date": work_date,
        "customer_name": customer_name,
        "worker_name": worker_name,
    }

    files: List[Tuple[str, Tuple[str, io.BytesIO, str]]] = []
    for item in uploaded_files:
        files.append(uploaded_to_file_tuple(item))
    if captured_photo:
        files.append(camera_bytes_to_file_tuple(captured_photo))
    if recorded_audio is not None:
        files.append(recorded_audio_to_file_tuple(recorded_audio))

    headers = {
        "x-api-key": API_KEY,
        "accept": "application/json",
        "Expect": "",
    }

    log_stdout(
        {
            "endpoint": endpoint,
            "timeout_sec": TIMEOUT_SEC,
            "equipment_model": equipment_model,
            "work_memo_len": len(work_memo),
            "file_count": len(files),
        },
        prefix="--- REQUEST DEBUG ---",
    )

    resp = requests.post(
        endpoint,
        headers=headers,
        data=data,
        files=files if files else None,
        timeout=TIMEOUT_SEC,
    )

    log_stdout(
        {
            "http_status": resp.status_code,
            "content_type": resp.headers.get("content-type", ""),
        },
        prefix="--- RESPONSE DEBUG ---",
    )

    if not resp.ok:
        log_stdout(resp.text, prefix="--- ERROR BODY ---")
        raise requests.HTTPError(f"HTTP {resp.status_code}", response=resp)

    try:
        parsed = resp.json()
    except Exception:
        log_stdout(resp.text, prefix="--- NON-JSON BODY ---")
        raise

    log_stdout(parsed, prefix="--- RESULT JSON RAW ---")
    normalized = normalize_result(parsed)
    log_stdout(
        {
            "normalized_keys": list(normalized.keys()) if isinstance(normalized, dict) else [],
            "reference_count": len(safe_list(normalized.get("references"))) if isinstance(normalized, dict) else 0,
        },
        prefix="--- RESULT NORMALIZED SUMMARY ---",
    )
    return normalized





def log_word_generation_error(ex: Exception, result: Dict[str, Any]) -> str:
    """Word生成エラーを標準出力へ詳細出力し、画面表示用メッセージを返す。"""
    tb = traceback.format_exc()
    diagnostics = {
        "error_type": type(ex).__name__,
        "error_message": str(ex),
        "python_executable": sys.executable,
        "docx_available": DOCX_AVAILABLE,
        "docx_import_error": DOCX_IMPORT_ERROR,
        "document_object_is_none": Document is None,
        "pt_object_is_none": Pt is None,
        "result_top_level_keys": list(result.keys()) if isinstance(result, dict) else [],
    }
    log_stdout(diagnostics, prefix="--- WORD GENERATION ERROR DIAGNOSTICS ---")
    log_stdout(tb, prefix="--- WORD GENERATION ERROR TRACEBACK ---")
    return (
        "Wordファイルの生成に失敗しました。"
        f" error_type={type(ex).__name__}, message={str(ex)}"
    )


def render_docx_import_error() -> None:
    """python-docx import失敗時の診断情報を表示する。"""
    st.warning(
        "Wordダウンロード機能を使うには、Streamlitを起動しているPython環境で "
        "`from docx import Document` が成功する必要があります。"
    )
    with st.expander("python-docx import 診断情報"):
        st.write("Streamlit実行中のPython:", sys.executable)
        st.write("DOCX_AVAILABLE:", DOCX_AVAILABLE)
        st.write("DOCX_IMPORT_ERROR:", DOCX_IMPORT_ERROR or "なし")
        try:
            import docx
            st.write("import docx 直接確認:", "OK")
            st.write("docx module path:", getattr(docx, "__file__", "不明"))
        except Exception as direct_import_error:
            st.write("import docx 直接確認:", repr(direct_import_error))
        st.code(
            """python -c "import sys; print(sys.executable); from docx import Document; print('python-docx OK')"
python -m pip show python-docx
python -m streamlit run app.py""",
            language="bash",
        )

# -----------------------------------------------------------
# Wordファイル生成
# -----------------------------------------------------------
def _docx_add_heading(document: Any, text: str, level: int = 1) -> None:
    """Word文書へ見出しを追加する。"""
    document.add_heading(text or "", level=level)


def _docx_add_paragraph(document: Any, label: str, value: Any) -> None:
    """ラベル付き段落を追加する。"""
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run("" if value is None else str(value))


def _docx_add_bullets(document: Any, items: List[Any]) -> None:
    """箇条書きを追加する。"""
    for item in items:
        if isinstance(item, dict):
            text = " / ".join(
                f"{k}: {v}" for k, v in item.items()
                if v not in (None, "", [])
            )
        else:
            text = str(item)
        document.add_paragraph(text, style="List Bullet")


def create_work_report_docx_bytes(result: Dict[str, Any]) -> bytes:
    """エージェントのレスポンスから作業報告書Wordファイルのバイト列を生成する。"""
    if not DOCX_AVAILABLE or Document is None or Pt is None:
        raise RuntimeError(
            "Streamlit起動時に python-docx をimportできませんでした。"
            f" import失敗内容: {DOCX_IMPORT_ERROR}"
        )

    report = safe_dict(result.get("report"))
    equipment = safe_dict(report.get("equipment"))
    labor = safe_dict(report.get("labor"))
    completion = safe_dict(report.get("completion_status"))

    document = Document()

    styles = document.styles
    styles["Normal"].font.name = "Yu Gothic"
    styles["Normal"].font.size = Pt(10.5)

    title = report.get("report_title") or "作業報告書"
    document.add_heading(title, level=0)

    _docx_add_heading(document, "1. 基本情報", 1)
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    def add_row(label: str, value: Any) -> None:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = "" if value is None else str(value)

    add_row("作業日", report.get("work_date", ""))
    add_row("顧客名", report.get("customer_name", ""))
    add_row("作業者名", report.get("worker_name", ""))
    add_row("設備型番", equipment.get("model", ""))
    add_row("設備名", equipment.get("name", ""))
    add_row("設置場所", equipment.get("location", ""))
    add_row("作業分類", report.get("work_classification", ""))

    _docx_add_heading(document, "2. 作業概要", 1)
    _docx_add_paragraph(document, "受付内容・症状", report.get("reception_summary", ""))
    _docx_add_paragraph(document, "音声メモ要約", report.get("transcript_summary", ""))
    _docx_add_paragraph(document, "作業メモ要約", report.get("work_memo_summary", ""))
    _docx_add_paragraph(document, "作業前の状態", report.get("before_condition", ""))

    _docx_add_heading(document, "3. 実施作業", 1)
    performed_work = safe_list(report.get("performed_work"))
    if performed_work:
        work_table = document.add_table(rows=1, cols=4)
        work_table.style = "Table Grid"
        headers = ["工程", "作業アクション", "作業詳細", "作業結果"]
        for i, h in enumerate(headers):
            work_table.rows[0].cells[i].text = h
        for idx, item in enumerate(performed_work, 1):
            row = safe_dict(item)
            cells = work_table.add_row().cells
            cells[0].text = str(row.get("step_no") or idx)
            cells[1].text = str(row.get("action", ""))
            cells[2].text = str(row.get("detail", ""))
            cells[3].text = str(row.get("result", ""))
    else:
        document.add_paragraph("実施作業工程はありません。")

    _docx_add_heading(document, "4. 使用部品・工数", 1)
    _docx_add_paragraph(document, "作業時間", f"{labor.get('hours', 0)} 時間")
    _docx_add_paragraph(document, "工数の算出根拠", labor.get("basis", ""))
    _docx_add_paragraph(document, "工数信頼度", labor.get("confidence", ""))

    used_parts = safe_list(report.get("used_parts"))
    if used_parts:
        parts_table = document.add_table(rows=1, cols=5)
        parts_table.style = "Table Grid"
        headers = ["部品名", "部品型番", "数量", "単位", "使用理由"]
        for i, h in enumerate(headers):
            parts_table.rows[0].cells[i].text = h
        for item in used_parts:
            row = safe_dict(item)
            cells = parts_table.add_row().cells
            cells[0].text = str(row.get("part_name", ""))
            cells[1].text = str(row.get("part_model", ""))
            cells[2].text = str(row.get("quantity", 0))
            cells[3].text = str(row.get("unit", ""))
            cells[4].text = str(row.get("reason", ""))
    else:
        document.add_paragraph("使用部品はありません。")

    _docx_add_heading(document, "5. 測定値・確認結果", 1)
    measurements = safe_list(report.get("measurements_and_checks"))
    if measurements:
        m_table = document.add_table(rows=1, cols=4)
        m_table.style = "Table Grid"
        headers = ["項目", "測定値・状態", "判定基準", "判定結果"]
        for i, h in enumerate(headers):
            m_table.rows[0].cells[i].text = h
        for item in measurements:
            row = safe_dict(item)
            cells = m_table.add_row().cells
            cells[0].text = str(row.get("item", ""))
            cells[1].text = str(row.get("value", ""))
            cells[2].text = str(row.get("criterion", ""))
            cells[3].text = str(row.get("result", ""))
    else:
        document.add_paragraph("測定値・点検項目はありません。")

    photo_findings = safe_list(report.get("photo_findings"))
    if photo_findings:
        _docx_add_heading(document, "6. 写真からの確認事項", 1)
        for item in photo_findings:
            row = safe_dict(item)
            _docx_add_paragraph(
                document,
                row.get("file_name", "写真"),
                f"{row.get('finding', '')}（信頼度: {row.get('confidence', '')}）",
            )

    _docx_add_heading(document, "7. 完了状況・申し送り", 1)
    _docx_add_paragraph(document, "完了ステータス", completion.get("status", ""))
    _docx_add_paragraph(document, "理由", completion.get("reason", ""))
    _docx_add_paragraph(document, "顧客への申し送り事項", report.get("customer_notes", ""))
    next_actions = safe_list(report.get("next_recommended_actions"))
    if next_actions:
        document.add_paragraph("次回推奨アクション:")
        _docx_add_bullets(document, next_actions)

    _docx_add_heading(document, "8. 請求項目候補", 1)
    billing = safe_list(result.get("billing_candidates"))
    if billing:
        b_table = document.add_table(rows=1, cols=5)
        b_table.style = "Table Grid"
        headers = ["項目名", "数量", "単位", "理由", "信頼度"]
        for i, h in enumerate(headers):
            b_table.rows[0].cells[i].text = h
        for item in billing:
            row = safe_dict(item)
            cells = b_table.add_row().cells
            cells[0].text = str(row.get("item_name", ""))
            cells[1].text = str(row.get("quantity", 0))
            cells[2].text = str(row.get("unit", ""))
            cells[3].text = str(row.get("reason", ""))
            cells[4].text = str(row.get("confidence", ""))
    else:
        document.add_paragraph("請求項目候補はありません。")

    _docx_add_heading(document, "9. 不足情報・不明瞭箇所", 1)
    missing = safe_list(result.get("missing_information"))
    if missing:
        document.add_paragraph("不足情報:")
        _docx_add_bullets(document, missing)
    else:
        document.add_paragraph("不足情報はありません。")

    uncertain = safe_list(result.get("uncertain_transcript_segments"))
    if uncertain:
        document.add_paragraph("音声文字起こしの不明瞭箇所:")
        _docx_add_bullets(document, uncertain)
    else:
        document.add_paragraph("音声文字起こしの不明瞭箇所はありません。")

    _docx_add_heading(document, "10. データソース根拠", 1)
    refs = safe_list(result.get("references"))
    if refs:
        for item in refs:
            row = safe_dict(item)
            _docx_add_paragraph(
                document,
                row.get("source", "参照元"),
                f"{row.get('used_for', '')}: {row.get('detail', '')}",
            )
    else:
        document.add_paragraph("データソース参照情報はありません。")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def build_docx_file_name(result: Dict[str, Any]) -> str:
    """作業報告書のダウンロードファイル名を作成する。"""
    report = safe_dict(result.get("report"))
    equipment = safe_dict(report.get("equipment"))
    work_date = str(report.get("work_date") or "work_date_unknown").replace("/", "-")
    model = str(equipment.get("model") or "equipment_unknown")
    return f"作業報告書_{work_date}_{model}.docx"


# -----------------------------------------------------------
# 表示部品
# -----------------------------------------------------------
def render_empty(message: str, icon: str = "📋") -> None:
    st.markdown(
        f'<div class="empty-state"><div class="empty-icon">{esc(icon)}</div>'
        f'<div>{esc(message)}</div></div>',
        unsafe_allow_html=True,
    )


def render_kv(label: str, value: Any) -> None:
    st.markdown(
        '<div class="kv-card">'
        f'<div class="kv-label">{esc(label)}</div>'
        f'<div class="kv-value">{esc(value)}</div>'
        '</div>',
        unsafe_allow_html=True,
    )


def render_status_badge(label: str, status: str = "") -> None:
    cls = "status-badge"
    if status in ("ok", "completed", "high"):
        cls += " status-ok"
    elif status in ("ng", "follow_up_required", "low"):
        cls += " status-ng"
    elif status in ("unknown", "temporary_recovery", "medium"):
        cls += " status-warn"
    st.markdown(f'<span class="{cls}">{esc(label)}</span>', unsafe_allow_html=True)


def render_report_info_table(rows: List[Tuple[str, Any, str, Any]]) -> None:
    """報告書タブの基本情報を2項目横並びの表で表示する。"""
    body = ""
    for left_label, left_value, right_label, right_value in rows:
        body += (
            "<tr>"
            f"<th>{esc(left_label)}</th>"
            f"<td>{esc(left_value)}</td>"
            f"<th>{esc(right_label)}</th>"
            f"<td>{esc(right_value)}</td>"
            "</tr>"
        )
    st.markdown(
        f'<table class="report-info-table"><tbody>{body}</tbody></table>',
        unsafe_allow_html=True,
    )


def render_report_text_table(rows: List[Tuple[str, Any]]) -> None:
    """報告書タブの長文項目を大きな表で表示する。"""
    body = ""
    for label, value in rows:
        body += (
            "<tr>"
            f"<th>{esc(label)}</th>"
            f"<td>{esc(value)}</td>"
            "</tr>"
        )
    st.markdown(
        f'<table class="report-text-table"><tbody>{body}</tbody></table>',
        unsafe_allow_html=True,
    )


def render_report_section_heading(title: str) -> None:
    """報告書タブ内の小見出しを表示する。"""
    st.markdown(
        f'<div class="report-section-heading">{esc(title)}</div>',
        unsafe_allow_html=True,
    )


def render_report_overview(report: Dict[str, Any]) -> None:
    equipment = safe_dict(report.get("equipment"))
    completion = safe_dict(report.get("completion_status"))

    equipment_name = equipment.get("name", "")
    equipment_location = equipment.get("location", "")
    if equipment_name and equipment_location:
        equipment_name_location = f"{equipment_name} / {equipment_location}"
    else:
        equipment_name_location = equipment_name or equipment_location

    render_report_section_heading("基本情報")
    render_report_info_table([
        ("報告書タイトル", report.get("report_title", ""), "作業分類", report.get("work_classification", "")),
        ("作業日", report.get("work_date", ""), "完了ステータス", completion.get("status", "")),
        ("顧客名", report.get("customer_name", ""), "作業者名", report.get("worker_name", "")),
        ("設備型番", equipment.get("model", ""), "設備名 / 設置場所", equipment_name_location),
    ])

    if completion.get("reason"):
        render_report_text_table([
            ("完了ステータスの理由", completion.get("reason", "")),
        ])

    render_report_section_heading("作業内容")
    render_report_text_table([
        ("受付内容・症状の要約", report.get("reception_summary", "")),
        ("音声メモの文字起こし要約", report.get("transcript_summary", "")),
        ("作業メモの要約", report.get("work_memo_summary", "")),
        ("作業前の状態", report.get("before_condition", "")),
        ("顧客への申し送り事項", report.get("customer_notes", "")),
    ])

    actions = safe_list(report.get("next_recommended_actions"))
    if actions:
        render_report_section_heading("次回推奨アクション")
        chips = ""
        for idx, action in enumerate(actions, 1):
            chips += f'<div class="chip"><span class="chip-num">{idx}</span>{esc(action)}</div>'
        st.markdown(
            f'<div class="chip-grid">{chips}</div>',
            unsafe_allow_html=True,
        )

def render_performed_work(items: List[Any]) -> None:
    if not items:
        render_empty("実施作業工程はありません", "🔧")
        return
    for idx, item in enumerate(items, 1):
        row = safe_dict(item)
        step_no = row.get("step_no") or idx
        st.markdown(
            '<div class="step-card">'
            f'<div class="step-header">STEP {esc(step_no)} / {esc(row.get("action", ""))}</div>'
            '<div class="step-body">'
            f'<p><b>作業詳細:</b> {esc(row.get("detail", ""))}</p>'
            f'<p><b>作業結果:</b> {esc(row.get("result", ""))}</p>'
            '</div></div>',
            unsafe_allow_html=True,
        )


def render_parts_and_labor(report: Dict[str, Any]) -> None:
    labor = safe_dict(report.get("labor"))
    cols = st.columns([1, 2])
    with cols[0]:
        st.markdown(
            '<div class="metric-card">'
            f'<div class="metric-value">{esc(labor.get("hours", 0))}</div>'
            '<div class="metric-label">作業時間（時間）</div>'
            '</div>',
            unsafe_allow_html=True,
        )
        if labor.get("confidence"):
            render_status_badge(f"信頼度: {labor.get('confidence')}", labor.get("confidence", ""))
    with cols[1]:
        render_kv("工数の算出根拠", labor.get("basis", ""))

    parts = safe_list(report.get("used_parts"))
    st.markdown("#### 使用部品")
    if not parts:
        render_empty("使用部品はありません", "📦")
        return
    for part in parts:
        row = safe_dict(part)
        render_kv(
            f"{row.get('part_name', '')} {row.get('part_model', '')}".strip() or "部品",
            f"数量: {row.get('quantity', 0)} {row.get('unit', '')}\n理由: {row.get('reason', '')}".strip(),
        )


def render_measurements(items: List[Any]) -> None:
    if not items:
        render_empty("測定値・点検項目はありません", "📏")
        return
    for item in items:
        row = safe_dict(item)
        render_status_badge(f"判定: {row.get('result', '')}", row.get("result", ""))
        render_kv(
            row.get("item", "点検項目"),
            f"測定値・状態: {row.get('value', '')}\n判定基準: {row.get('criterion', '')}",
        )


def render_billing(items: List[Any]) -> None:
    if not items:
        render_empty("請求項目候補はありません", "💴")
        return
    st.caption("表示内容は請求確定ではなく、作業報告書から抽出した候補です。")
    for item in items:
        row = safe_dict(item)
        render_status_badge(f"{row.get('item_type', '')} / 信頼度: {row.get('confidence', '')}", row.get("confidence", ""))
        render_kv(
            row.get("item_name", "請求項目候補"),
            f"数量: {row.get('quantity', 0)} {row.get('unit', '')}\n理由: {row.get('reason', '')}",
        )


def render_file_observations(items: List[Any]) -> None:
    if not items:
        render_empty("添付ファイルの観察事項はありません", "🖼️")
        return
    for item in items:
        row = safe_dict(item)
        facts = "\n".join(f"・{x}" for x in safe_list(row.get("observed_facts")))
        limits = "\n".join(f"・{x}" for x in safe_list(row.get("limitations")))
        render_kv(
            f"{row.get('file_name', '')} ({row.get('file_type', '')})",
            f"観察事項:\n{facts}\n\n制限事項:\n{limits}".strip(),
        )


def render_missing(items: List[Any]) -> None:
    if not items:
        render_empty("不足情報はありません", "✅")
        return
    for item in items:
        row = safe_dict(item)
        must = "確定に必須" if row.get("required_to_finalize") else "任意確認"
        render_status_badge(must, "ng" if row.get("required_to_finalize") else "unknown")
        render_kv(row.get("field", "不足項目"), row.get("reason", ""))


def render_uncertain_segments(items: List[Any]) -> None:
    if not items:
        render_empty("音声文字起こしの不明瞭箇所はありません", "🎙️")
        return
    for item in items:
        row = safe_dict(item)
        render_kv(row.get("text", "不明瞭箇所"), row.get("reason", ""))


def render_references(items: List[Any]) -> None:
    if not items:
        render_empty("データソース参照情報はありません", "📚")
        return
    for item in items:
        row = safe_dict(item)
        render_status_badge(row.get("used_for", ""), "")
        render_kv(row.get("source", "参照元"), row.get("detail", ""))


# ===========================================================
# メイン UI
# ===========================================================
st.markdown(
    '<div class="hero-banner">'
    '<h1>📝 作業報告書自動生成アプリ</h1>'
    '<p>作業写真・音声メモ・作業内容メモ・設備型番をエージェントに送信し、標準フォーマット作業報告書と請求項目候補を生成します。</p>'
    '</div>',
    unsafe_allow_html=True,
)

if not API_BASE_URL or not API_KEY:
    st.error(
        "環境変数が不足しています。\n\n"
        "- API_BASE_URL（例: https://agent.example.com/xxxx）\n"
        "- API_KEY（x-api-key の値）\n"
        "- API_TIMEOUT_SEC（任意。未指定時は 600 秒）"
    )
    st.stop()

for key, default in [
    ("last_result", None),
    ("captured_photo", None),
    ("camera_mode", "idle"),
]:
    if key not in st.session_state:
        st.session_state[key] = default

left, right = st.columns([1, 1.25], gap="large")

with left:
    st.markdown(
        '<div class="section-card"><div class="section-title"><span class="icon">📝</span>入力情報</div>',
        unsafe_allow_html=True,
    )

    equipment_model = st.selectbox(
        "設備型番（必須）",
        options=["選択してください"] + EQUIPMENT_OPTIONS,
        index=0,
    )

    col1, col2 = st.columns(2)
    with col1:
        selected_work_date = st.date_input(
            "作業日",
            value=date.today(),
            format="YYYY-MM-DD",
        )
        work_date = selected_work_date.isoformat()
    with col2:
        worker_name = st.text_input("作業者名", placeholder="例: 山田太郎")

    customer_name = st.text_input("顧客名", placeholder="例: 株式会社サンプル")

    work_memo = st.text_area(
        "作業内容メモ",
        height=160,
        placeholder="例: 搬送コンベアのベルトが右側に寄る。ガイドローラー2個を交換し、張力調整。作業時間2時間。試運転30分で再発なし。",
    )

    uploaded_files = st.file_uploader(
        "添付ファイル（作業写真・音声メモ・作業メモなど / 複数可）",
        type=None,
        accept_multiple_files=True,
        help="未添付でも実行できます。",
    )

    st.markdown(
        '<div style="font-size:0.9rem;font-weight:700;margin:0.75rem 0 0.25rem;">'
        '🎙️ 音声メモ録音'
        '</div>',
        unsafe_allow_html=True,
    )
    recorded_audio = st.audio_input(
        "録音",
        key="recorded_voice_memo",
        help="マイクボタンをクリックして録音できます。録音後、WAVファイルとしてAPIへ送信します。",
    )
    if recorded_audio is not None:
        st.caption("録音済み: 音声ファイルとしてエージェントに送信されます。")

    if back_camera_input is not None:
        cam = st.session_state["camera_mode"]
        if cam == "idle":
            if st.button("📷 カメラで作業写真を撮る", use_container_width=True):
                st.session_state["camera_mode"] = "camera"
                st.rerun()
        elif cam == "camera":
            photo = back_camera_input(key="rear_camera")
            if photo is not None:
                st.session_state["captured_photo"] = photo.getvalue()
                st.session_state["camera_mode"] = "preview"
                st.rerun()
        elif cam == "preview":
            st.image(st.session_state["captured_photo"], use_container_width=True)
            if st.button("📷 撮り直す", use_container_width=True):
                st.session_state["captured_photo"] = None
                st.session_state["camera_mode"] = "camera"
                st.rerun()
    else:
        st.caption("カメラ入力を使う場合は streamlit-back-camera-input をインストールしてください。")

    run = st.button("🚀 作業報告書を生成", type="primary", use_container_width=True)
    st.markdown("</div>", unsafe_allow_html=True)

    if run:
        errors: List[str] = []
        if equipment_model == "選択してください":
            errors.append("設備型番を選択してください。")
        if (
            not work_memo.strip()
            and not uploaded_files
            and not st.session_state.get("captured_photo")
            and recorded_audio is None
        ):
            errors.append("作業内容メモ、添付ファイル、カメラ写真、録音音声のいずれかを入力してください。")

        if errors:
            for error in errors:
                st.error(error)
        else:
            try:
                with st.spinner("エージェントを呼び出しています..."):
                    st.session_state["last_result"] = call_api(
                        equipment_model=equipment_model,
                        work_memo=work_memo.strip(),
                        work_date=work_date,
                        customer_name=customer_name.strip(),
                        worker_name=worker_name.strip(),
                        uploaded_files=uploaded_files or [],
                        captured_photo=st.session_state.get("captured_photo"),
                        recorded_audio=recorded_audio,
                    )
            except requests.Timeout:
                st.error("タイムアウトしました。API_TIMEOUT_SEC を延ばすか、サーバー側の処理状況を確認してください。")
            except requests.HTTPError as ex:
                status = getattr(getattr(ex, "response", None), "status_code", "")
                st.error(f"エージェントでHTTPエラーが発生しました。status={status}。詳細は標準出力ログを確認してください。")
            except Exception as ex:
                log_stdout(str(ex), prefix="--- UNEXPECTED ERROR ---")
                st.error("予期せぬエラーが発生しました。詳細は標準出力ログを確認してください。")

with right:
    st.markdown(
        '<div class="section-card">'
        '<div class="section-title"><span class="icon">📄</span>生成結果</div>',
        unsafe_allow_html=True,
    )

    result = st.session_state.get("last_result")
    if not result:
        st.markdown(
            '<div class="empty-state"><div class="empty-icon">📋</div>'
            '<div>左側で入力して「作業報告書を生成」を押すと、ここに結果が表示されます。</div>'
            '</div>',
            unsafe_allow_html=True,
        )
    else:
        schema_ok, missing_keys = validate_result_schema(result)
        if not schema_ok:
            log_schema_mismatch(result, missing_keys)
            st.error(
                "応答スキーマが想定と異なります。標準出力ログに不足キー・実際のキー・レスポンス全文を出力しました。"
            )
            with st.expander("画面上で不足キーを確認する"):
                st.write("不足キー:", missing_keys)
                st.write("実際のキー:", list(result.keys()) if isinstance(result, dict) else [])
            st.stop()

        if DOCX_AVAILABLE:
            try:
                docx_bytes = create_work_report_docx_bytes(result)
                st.download_button(
                    label="📥 作業報告書(Word)をダウンロード",
                    data=docx_bytes,
                    file_name=build_docx_file_name(result),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as ex:
                message = log_word_generation_error(ex, result)
                st.error(message)
                with st.expander("Word生成エラーの詳細"):
                    st.write("Streamlit実行中のPython:", sys.executable)
                    st.write("DOCX_AVAILABLE:", DOCX_AVAILABLE)
                    st.write("DOCX_IMPORT_ERROR:", DOCX_IMPORT_ERROR or "なし")
                    st.code(traceback.format_exc(), language="text")
        else:
            render_docx_import_error()


        report = safe_dict(result.get("report"))

        tabs = st.tabs([
            "📄 報告書",
            "🔧 実施作業",
            "📦 部品・工数",
            "📏 測定・確認",
            "💴 請求候補",
            "🖼️ 添付解析",
            "⚠️ 不足・不明瞭",
            "📚 データソース根拠",
        ])

        with tabs[0]:
            render_report_overview(report)

        with tabs[1]:
            render_performed_work(safe_list(report.get("performed_work")))

        with tabs[2]:
            render_parts_and_labor(report)

        with tabs[3]:
            render_measurements(safe_list(report.get("measurements_and_checks")))
            photo_findings = safe_list(report.get("photo_findings"))
            if photo_findings:
                st.markdown("#### 写真からの確認事項")
                for item in photo_findings:
                    row = safe_dict(item)
                    render_status_badge(f"信頼度: {row.get('confidence', '')}", row.get("confidence", ""))
                    render_kv(row.get("file_name", "写真"), row.get("finding", ""))

        with tabs[4]:
            render_billing(safe_list(result.get("billing_candidates")))

        with tabs[5]:
            render_file_observations(safe_list(result.get("attached_file_observations")))

        with tabs[6]:
            st.markdown("#### 不足情報")
            render_missing(safe_list(result.get("missing_information")))
            st.markdown("#### 音声文字起こしの不明瞭箇所")
            render_uncertain_segments(safe_list(result.get("uncertain_transcript_segments")))

        with tabs[7]:
            render_references(safe_list(result.get("references")))

    st.markdown("</div>", unsafe_allow_html=True)
